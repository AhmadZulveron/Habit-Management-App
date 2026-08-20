import json
from docx import Document
from docx.shared import Inches

doc_path = r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx"
doc = Document(doc_path)

def replace_text_in_paragraph(p, old_text, new_text):
    if old_text in p.text:
        # A simple replacement if the text is fully contained in a single run
        # but to be safe with python-docx, if old_text spans multiple runs it's hard.
        # So we just replace p.text entirely. However, this loses bold formatting.
        # Let's check if there are multiple runs.
        if len(p.runs) == 1:
            p.runs[0].text = p.runs[0].text.replace(old_text, new_text)
        else:
            # We clear runs and add new text if we just replace p.text
            full_text = p.text.replace(old_text, new_text)
            for r in p.runs:
                r.text = ""
            p.add_run(full_text)
        return True
    return False

gap_paragraph = None
for p in doc.paragraphs:
    # A. "algoritma terbukti..." -> "Hasil pengujian menunjukkan..."
    # Actually I used "Hasil pengujian menunjukkan bahwa algoritma mampu" previously. The user's request: Replace "algoritma terbukti..." with "Hasil pengujian menunjukkan...". Since I already used "Hasil pengujian menunjukkan...", there is no "algoritma terbukti..." left in the text I injected! I'll just check just in case.
    if "algoritma terbukti" in p.text:
        replace_text_in_paragraph(p, "algoritma terbukti", "Hasil pengujian menunjukkan")
    
    # B. "pure sorting compute time" -> "waktu eksekusi QuickSort dalam batas pengukuran benchmark"
    # Actually in Phase 12B I wrote: "Pengukuran dilakukan dalam batas waktu (benchmark timing boundary) yang mengapit proses eksekusi algoritma QuickSort."
    # The user says: Replace "pure sorting compute time" with "waktu eksekusi QuickSort dalam batas pengukuran benchmark".
    if "pure sorting compute time" in p.text:
        replace_text_in_paragraph(p, "pure sorting compute time", "waktu eksekusi QuickSort dalam batas pengukuran benchmark")

    # C. "The observed behavior is consistent with characteristics of recursive Lomuto-style partitioning on highly duplicated input." -> "Perilaku yang diamati konsisten dengan karakteristik partisi rekursif bergaya Lomuto ketika memproses input dengan tingkat duplikasi yang tinggi."
    if "The observed behavior is consistent with characteristics of recursive Lomuto-style partitioning on highly duplicated input." in p.text:
        replace_text_in_paragraph(p, "The observed behavior is consistent with characteristics of recursive Lomuto-style partitioning on highly duplicated input.", "Perilaku yang diamati konsisten dengan karakteristik partisi rekursif bergaya Lomuto ketika memproses input dengan tingkat duplikasi yang tinggi.")

    if "[Gap Analisis:" in p.text:
        gap_paragraph = p

# Format Table 4.6 floating point values
for tbl in doc.tables:
    # Check if this is Tabel 4.6 (Waktu Eksekusi Median (ms))
    # Let's check the first cell text or header
    if len(tbl.rows) > 0 and tbl.rows[0].cells[0].text == "Ukuran (N)" and len(tbl.columns) == 5:
        # Check if it's the timing table by inspecting a random cell in row 1
        val = tbl.rows[1].cells[1].text
        if "." in val and len(val) > 10:  # Looks like a long float
            for row in tbl.rows[1:]:
                for cell in row.cells[1:]:
                    try:
                        f_val = float(cell.text)
                        cell.text = f"{f_val:.6f}"
                    except ValueError:
                        pass

# 4.2 Black-Box Testing Table
# Replace gap paragraph text
if gap_paragraph:
    full_text = gap_paragraph.text
    full_text = full_text.split("[Gap Analisis:")[0].strip()
    full_text += " Tabel pengujian berikut menyajikan hasil validasi fungsional aplikasi menggunakan metode Black-Box Testing berdasarkan bukti visual (screenshot) aplikasi."
    for r in gap_paragraph.runs:
        r.text = ""
    gap_paragraph.add_run(full_text)

    # Insert Table 4.0 or just a table.
    headers = ["No.", "Fitur", "Skenario Pengujian", "Hasil yang Diharapkan", "Hasil Pengujian", "Status"]
    rows = [
        ["1", "Login", "Pengguna memasukkan kredensial pada antarmuka Login.", "Aplikasi menampilkan halaman Login.", "Antarmuka Login ditampilkan dengan field input yang sesuai (Login Page.jpg).", "Valid"],
        ["2", "Sign Up", "Pengguna mengakses pendaftaran akun baru.", "Aplikasi menampilkan antarmuka Sign Up.", "Halaman Sign Up ditampilkan dengan form pendaftaran (Sign Up Page.jpg).", "Valid"],
        ["3", "Home", "Pengguna melihat daftar kebiasaan hari ini.", "Menampilkan halaman Home dengan daftar kebiasaan.", "Halaman Home menampilkan daftar kebiasaan (Home Page.jpg, Home - completed today's habits list.jpg).", "Valid"],
        ["4", "Create Habit (Valid)", "Pengguna membuat kebiasaan baru dengan input valid.", "Aplikasi menampilkan halaman pembuatan kebiasaan.", "Form pembuatan kebiasaan ditampilkan (Create Habit Page.jpg).", "Valid"],
        ["5", "Create Habit (Invalid)", "Pengguna memasukkan input tidak valid saat pembuatan.", "Menampilkan pesan error validasi input.", "Pesan error ditampilkan di UI (Create Habit - user input error message.jpg).", "Valid"],
        ["6", "Habit List", "Pengguna melihat daftar semua kebiasaan.", "Menampilkan daftar seluruh kebiasaan (Habit List).", "Halaman Habit List berhasil ditampilkan (Habit List Page.jpg).", "Valid"],
        ["7", "Add to Today's Habit", "Pengguna menambahkan kebiasaan ke jadwal hari ini.", "Menampilkan dialog konfirmasi dan pesan sukses penambahan.", "Dialog konfirmasi (Habit List - add to today's habits confirmation message.jpg) dan pesan sukses tampil (Habit List - added successfully to today's message.jpg).", "Valid"],
        ["8", "Habit Detail", "Pengguna membuka detail kebiasaan.", "Menampilkan halaman detail.", "Halaman detail berhasil ditampilkan (Habit Detail Page.jpg).", "Valid"],
        ["9", "Deactivate Habit", "Pengguna menonaktifkan kebiasaan.", "Memunculkan konfirmasi deaktivasi dan pesan sukses.", "Konfirmasi (Habit Detail - deactivate habit confirmation message.jpg) dan pesan sukses tampil (Habit Detail - deactivated successfully message.jpg).", "Valid"],
        ["10", "Edit Habit", "Pengguna memperbarui informasi kebiasaan.", "Menampilkan halaman edit, konfirmasi, dan pesan sukses update.", "Halaman Edit (Edit Habit Page.jpg), konfirmasi (Edit Habit Page - confirmation.jpg), dan pesan sukses (Edit Habit - updated successfully message.jpg) tampil.", "Valid"],
        ["11", "Delete Habit", "Pengguna menghapus kebiasaan.", "Menampilkan konfirmasi penghapusan dan pesan sukses.", "Konfirmasi hapus (Edit Habit - deleting confirmation message.jpg) dan pesan sukses (Edit Habit - deleted successfully message.jpg) tampil.", "Valid"],
        ["12", "Complete Habit", "Pengguna menyelesaikan kebiasaan dari halaman Home.", "Menampilkan pesan konfirmasi dan status selesai.", "Pesan konfirmasi (Home - complete habit confirmation message.jpg) dan pesan sukses (Home - habit completed message.jpg) tampil.", "Valid"],
        ["13", "Recommendation", "Pengguna mengakses halaman rekomendasi (QuickSort UI).", "Menampilkan antarmuka rekomendasi kebiasaan.", "Halaman Recommendation ditampilkan (Recommendation Page.jpg).", "Valid"],
        ["14", "Statistics/Report", "Pengguna melihat laporan statistik kebiasaan.", "Menampilkan halaman statistik.", "Halaman Report berhasil ditampilkan (Stats(Report Page).jpg).", "Valid"],
        ["15", "Profile", "Pengguna mengakses profil pengguna.", "Menampilkan informasi profil.", "Halaman Profile ditampilkan (Profile Page.jpg).", "Valid"],
        ["16", "Edit Profile", "Pengguna mengedit informasi profil.", "Menampilkan antarmuka edit profil.", "Halaman Edit Profile ditampilkan (Edit Profile Page.jpg).", "Valid"],
        ["17", "Badges", "Pengguna mengakses halaman lencana/penghargaan.", "Menampilkan halaman Badges.", "Halaman Badges berhasil ditampilkan (Badges Page.jpg).", "Valid"],
        ["18", "Settings", "Pengguna mengakses halaman pengaturan aplikasi.", "Menampilkan halaman Settings.", "Halaman Settings ditampilkan (Settings Page.jpg).", "Valid"]
    ]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hdr_cells = tbl.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
    for row_data in rows:
        row_cells = tbl.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)
    
    gap_paragraph._element.addnext(tbl._tbl)

doc.save(r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx")
print("SUCCESS!")

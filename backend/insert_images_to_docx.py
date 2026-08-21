from docx import Document
from docx.shared import Inches
import os

doc_path = r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx"
doc = Document(doc_path)
sym_dir = r"d:\flutterAhmadZulveron\Try\backend\symbols"

for table in doc.tables:
    for row in table.rows:
        if len(row.cells) >= 2:
            sym_name = row.cells[1].text.strip()
            cell = row.cells[0]
            txt = cell.text.strip()
            img_file = None
            
            if sym_name == "Actor" and "Actor Icon" in txt: img_file = "actor.png"
            elif sym_name == "Use Case" and "Oval" in txt: img_file = "use_case.png"
            elif sym_name == "Association": img_file = "association.png"
            
            elif sym_name == "Actor / Object": img_file = "object_box.png"
            elif sym_name == "Lifeline": img_file = "lifeline.png"
            elif sym_name == "Activation Box": img_file = "activation_box.png"
            elif sym_name == "Message": img_file = "sync_message.png"
            elif sym_name == "Return Message": img_file = "return_message.png"
            
            elif sym_name == "Initial Node": img_file = "initial_node.png"
            elif sym_name == "Activity": img_file = "activity.png"
            elif sym_name == "Control Flow": img_file = "control_flow.png"
            elif sym_name == "Decision Node": img_file = "decision.png"
            elif sym_name == "Final Node": img_file = "final_node.png"
            
            elif sym_name == "Terminator": img_file = "terminator.png"
            elif sym_name == "Process": img_file = "process.png"
            elif sym_name == "Decision": img_file = "decision.png"
            elif sym_name == "Input / Output": img_file = "io.png"
            elif sym_name == "Flow Line": img_file = "control_flow.png" # vertical arrow
            elif sym_name == "Predefined Process": img_file = "predefined.png"
            
            if img_file and "(" in txt and ")" in txt:
                full_img_path = os.path.join(sym_dir, img_file)
                if os.path.exists(full_img_path):
                    # Clear cell text safely
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = 1 # Center align
                    r = p.add_run()
                    # Determine width based on aspect ratio roughly
                    width = Inches(0.4)
                    if img_file in ["lifeline.png", "activation_box.png", "actor.png", "control_flow.png"]:
                        width = Inches(0.2)
                    r.add_picture(full_img_path, width=width)

doc.save(doc_path)
print("Images injected successfully.")

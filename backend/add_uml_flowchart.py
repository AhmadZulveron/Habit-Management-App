from docx import Document
from docx.shared import Pt
import sys

doc_path = r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx"
doc = Document(doc_path)

target_idx = -1
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip().upper()
    if "METODOLOGI" in txt and i > 200:
        target_idx = i
        break

if target_idx == -1:
    print("Could not find BAB III METODOLOGI body paragraph")
    sys.exit(1)

bab3_para = doc.paragraphs[target_idx]

def add_heading(text, level, is_bold=True):
    p = bab3_para.insert_paragraph_before(text)
    if is_bold:
        for run in p.runs:
            run.bold = True
    return p

def add_body(text):
    p = bab3_para.insert_paragraph_before(text)
    return p

def add_table(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    
    for row_data in rows:
        row_cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    bab3_para._p.addprevious(tbl._tbl)
    bab3_para.insert_paragraph_before("")

add_heading("2.8 Unified Modeling Language (UML)", level=2)

add_heading("2.8.1 Pengertian UML", level=3)
add_body("Unified Modeling Language (UML) merupakan bahasa pemodelan standar yang digunakan untuk memvisualisasikan, merancang, dan mendokumentasikan sistem perangkat lunak (Fowler, 2003).")

add_heading("2.8.2 Fungsi UML dalam Penelitian", level=3)
add_body("Dalam penelitian ini, UML berfungsi untuk memodelkan arsitektur perangkat lunak sebelum tahap implementasi dilakukan. UML memberikan representasi visual mengenai interaksi antara pengguna (aktor) dengan fitur-fitur aplikasi manajemen kebiasaan harian. Selain itu, UML digunakan untuk menggambarkan urutan komunikasi data antara aplikasi Flutter dan backend, serta memodelkan alur proses sistem secara keseluruhan agar mudah dipahami.")

add_heading("2.8.3 Use Case Diagram", level=3)
add_body("Use Case Diagram adalah diagram UML yang menggambarkan interaksi antara satu atau lebih aktor dengan fungsionalitas yang disediakan oleh sistem. Dalam penelitian ini, diagram ini digunakan untuk mendefinisikan batas ruang lingkup sistem dan fungsionalitas utama yang dapat dilakukan oleh pengguna, seperti mengelola habit dan melihat rekomendasi.")
add_body("Simbol Use Case Diagram yang digunakan dalam penelitian ini ditunjukkan pada tabel berikut:")
add_table(
    ["Simbol", "Nama Simbol", "Keterangan"],
    [
        ["(Actor Icon)", "Actor", "Entitas eksternal (pengguna) yang berinteraksi dengan sistem."],
        ["(Oval)", "Use Case", "Fungsionalitas atau layanan spesifik yang disediakan sistem."],
        ["(Solid Line)", "Association", "Garis penghubung interaksi antara aktor dengan use case."]
    ]
)

add_heading("2.8.4 Sequence Diagram", level=3)
add_body("Sequence Diagram adalah diagram UML yang memodelkan aliran pesan dan interaksi antar objek secara sekuensial berdasarkan urutan waktu. Diagram ini digunakan dalam penelitian untuk memperlihatkan bagaimana aplikasi, antarmuka pengguna, dan server (backend) saling bertukar data pada proses-proses spesifik seperti pengelolaan habit, penyelesaian habit, dan rekomendasi habit.")
add_body("Simbol Sequence Diagram yang digunakan dalam penelitian ini ditunjukkan pada tabel berikut:")
add_table(
    ["Simbol", "Nama Simbol", "Keterangan"],
    [
        ["(Actor/Object Box)", "Actor / Object", "Pengguna atau komponen sistem yang berpartisipasi dalam interaksi."],
        ["(Dashed Line)", "Lifeline", "Menggambarkan waktu hidup suatu objek selama proses berlangsung."],
        ["(Rectangle)", "Activation Box", "Masa aktif sebuah objek saat mengeksekusi operasi (Focus of Control)."],
        ["(Solid Arrow)", "Message", "Pesan atau pemanggilan metode antar objek secara sinkron."],
        ["(Dashed Arrow)", "Return Message", "Pesan balasan atau hasil operasi dari sebuah message."]
    ]
)

add_heading("2.8.5 Activity Diagram", level=3)
add_body("Activity Diagram adalah diagram UML yang memodelkan alur aktivitas atau proses kerja dari sebuah sistem secara dinamis. Dalam penelitian ini, Activity Diagram digunakan untuk merinci setiap langkah komputasi sistem, mulai dari interaksi pengguna hingga proses penyimpanan data di backend.")
add_body("Simbol Activity Diagram yang digunakan dalam penelitian ini ditunjukkan pada tabel berikut:")
add_table(
    ["Simbol", "Nama Simbol", "Keterangan"],
    [
        ["(Solid Circle)", "Initial Node", "Titik awal dimulainya sebuah aktivitas."],
        ["(Rounded Rectangle)", "Activity", "Tindakan atau komputasi yang dilakukan sistem."],
        ["(Solid Arrow)", "Control Flow", "Garis panah penunjuk arah eksekusi aktivitas."],
        ["(Diamond)", "Decision Node", "Percabangan kondisi atau evaluasi (Yes/No)."],
        ["(Bullseye)", "Final Node", "Titik akhir berhentinya seluruh aktivitas."]
    ]
)


add_heading("2.9 Flowchart", level=2)

add_heading("2.9.1 Pengertian Flowchart", level=3)
add_body("Flowchart adalah representasi grafis yang menggambarkan urutan proses, pengambilan keputusan, serta aliran input dan output dari sebuah sistem atau algoritma.")

add_heading("2.9.2 Fungsi Flowchart dalam Penelitian", level=3)
add_body("Flowchart digunakan dalam penelitian ini untuk memvisualisasikan dua proses utama, yaitu alur kerja sistem secara keseluruhan (Flowchart Sistem) dan alur logika dari algoritma yang diterapkan (Flowchart Algoritma QuickSort). Visualisasi ini mempermudah pemahaman mengenai urutan instruksi prosedural yang dieksekusi oleh sistem (Kendall & Kendall, 2011).")

add_heading("2.9.3 Simbol Flowchart", level=3)
add_body("Simbol Flowchart yang digunakan dalam penelitian ini ditunjukkan pada tabel berikut:")
add_table(
    ["Simbol", "Nama Simbol", "Keterangan"],
    [
        ["(Oval)", "Terminator", "Titik awal (Start) atau titik akhir (End) dari sebuah proses."],
        ["(Rectangle)", "Process", "Instruksi komputasi atau pengolahan data."],
        ["(Diamond)", "Decision", "Evaluasi kondisi (pilihan Ya atau Tidak)."],
        ["(Parallelogram)", "Input / Output", "Proses penerimaan data masuk atau penampilan hasil."],
        ["(Solid Arrow)", "Flow Line", "Garis panah penunjuk arah aliran proses."],
        ["(Double-lined Rectangle)", "Predefined Process", "Pemanggilan fungsi rekursi (digunakan pada Quicksort)."]
    ]
)

doc.save(r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx")
print("Successfully injected UML and Flowchart sections before BAB III.")

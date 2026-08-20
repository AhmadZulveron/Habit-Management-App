import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx"
doc = Document(doc_path)

def insert_text_after(paragraph, text, bold=False):
    new_p = doc.add_paragraph()
    run = new_p.add_run(text)
    if bold:
        run.bold = True
    paragraph._element.addnext(new_p._element)
    return new_p

def insert_table_after(paragraph, headers, rows_data):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hdr_cells = tbl.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
    for row_data in rows_data:
        row_cells = tbl.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)
    paragraph._element.addnext(tbl._tbl)
    return tbl

def insert_picture_after(paragraph, img_path):
    new_p = doc.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    paragraph._element.addnext(new_p._element)
    return new_p

p_hasil = None
p_sistem = None
p_algo = None

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt == "Hasil Implementasi":
        p_hasil = p
    elif txt == "Pengujian Sistem":
        p_sistem = p
    elif txt == "Pengujian Algoritma Quicksort":
        p_algo = p

if p_hasil:
    insert_text_after(p_hasil, "Implementasi sistem mencakup pembuatan pipeline rekomendasi. Rule Engine digunakan untuk memfilter kandidat, dilanjutkan dengan Scoring Engine untuk menghitung skor relevansi berdasarkan faktor pencapaian pengguna. Algoritma QuickSort diintegrasikan pada tahap akhir untuk mengurutkan kandidat rekomendasi secara menurun (descending) berdasarkan skor relevansi tersebut, sebelum data dikembalikan melalui API kepada pengguna.")

if p_sistem:
    insert_text_after(p_sistem, "Pengujian fungsional sistem secara keseluruhan (seperti manajemen kebiasaan dan fitur laporan) dilakukan menggunakan metode Black-Box Testing. [Gap Analisis: Hasil pengujian fungsional antarmuka dan interaksi pengguna belum disertakan secara spesifik dalam dokumentasi ini dan memerlukan pelengkapan pengujian sistem secara manual].")

def load_json(filename):
    with open(rf"d:\flutterAhmadZulveron\Try\backend\results\phase9\{filename}", 'r', encoding='utf-8') as f:
        return json.load(f)

if p_algo:
    curr_p = p_algo
    
    # 4.3.1
    curr_p = insert_text_after(curr_p, "4.3.1 Skenario Pengujian", bold=True)
    curr_p = insert_text_after(curr_p, "Pengujian performa QuickSort dieksekusi menggunakan konfigurasi ukuran dataset (N) yang bervariasi: 10, 50, 100, 500, 1.000, 5.000, dan 10.000 elemen. Data dibangkitkan secara deterministik menggunakan PRNG ke dalam empat distribusi input: Random, Ascending, Descending, dan Duplicate-Heavy. Setiap kombinasi ukuran dan distribusi dieksekusi sebanyak 50 kali pengulangan (repetitions). Nilai median digunakan sebagai parameter utama waktu eksekusi untuk mengurangi pengaruh lonjakan anomali akibat aktivitas runtime sistem operasi. Tabel 4.1 merangkum skenario tersebut.")
    
    table_data = load_json("final_benchmark_table.json")
    headers = ["N", "Distribusi Input", "Successful Runs", "Failed Runs", "Keterangan"]
    rows = []
    for item in table_data:
        rows.append([item["N"], item["Input Type"], item["Successful Runs"], item["Failed Runs"], item["Correctness"]])
    
    curr_p = insert_text_after(curr_p, "Tabel 4.1 Skenario Eksperimen Benchmark")
    tbl = insert_table_after(curr_p, headers, rows)
    new_p = doc.add_paragraph()
    tbl._element.addnext(new_p._element)
    curr_p = new_p

    # 4.3.2
    curr_p = insert_text_after(curr_p, "4.3.2 Kebenaran Hasil Pengurutan", bold=True)
    curr_p = insert_text_after(curr_p, "Evaluasi kebenaran memverifikasi bahwa pengurutan descending dan manipulasi keamanan objek (shallow copy) berjalan sesuai spesifikasi. Tercatat sebanyak 1.350 eksekusi yang selesai secara penuh, dan seluruh eksekusi tersebut (completed executions) berhasil melewati validasi kebenaran pengurutan (correctness validation). Kasus pengujian pada N=10.000 Duplicate-Heavy tidak selesai akibat RangeError dan murni diklasifikasikan sebagai keterbatasan runtime, bukan sebagai kegagalan validasi kebenaran (Tabel 4.2).")
    
    corr_data = load_json("correctness_table.json")
    curr_p = insert_text_after(curr_p, "Tabel 4.2 Validasi Kebenaran (Correctness)")
    corr_headers = ["Metric", "Result"]
    corr_rows = [[r["Metric"], r["Result"]] for r in corr_data]
    tbl2 = insert_table_after(curr_p, corr_headers, corr_rows)
    new_p = doc.add_paragraph()
    tbl2._element.addnext(new_p._element)
    curr_p = new_p

    curr_p = insert_text_after(curr_p, "Selain pengujian terkontrol (benchmark), dilakukan pula pengujian trace rekomendasi riil. Pengujian terpisah ini memverifikasi integrasi algoritma QuickSort secara langsung di dalam pipeline rekomendasi API (getRecommendations) yang terhubung dengan basis data MySQL di lingkungan pengujian (test environment). Hasil trace mengonfirmasi integrasi berhasil. Pengujian trace ini hanya mengukur kinerja QuickSort selama pemanggilan API, dan waktu eksekusinya tidak merepresentasikan total latensi respons API secara keseluruhan (Tabel 4.3).")
    
    trace_data = load_json("realistic_trace_table.json")
    curr_p = insert_text_after(curr_p, "Tabel 4.3 Trace Rekomendasi Riil")
    trace_hdr = ["Candidate Count", "Comparisons", "Swaps", "Sorting Time (ms)", "Correctness", "Input Source"]
    trace_rows = []
    for r in trace_data:
        trace_rows.append([r["Candidate Count"], r["Comparisons"], r["Swaps"], r["Sorting Time (ms)"], r["Correctness"], r["Input Source"]])
    tbl_trace = insert_table_after(curr_p, trace_hdr, trace_rows)
    new_p = doc.add_paragraph()
    tbl_trace._element.addnext(new_p._element)
    curr_p = new_p

    # 4.3.3
    curr_p = insert_text_after(curr_p, "4.3.3 Jumlah Perbandingan", bold=True)
    curr_p = insert_text_after(curr_p, "Jumlah perbandingan (comparisons) diukur secara aktual pada internal fungsi QuickSort. Pertumbuhan jumlah operasi perbandingan seiring bertambahnya ukuran data (N) diobservasi pada keempat distribusi input (Tabel 4.4 dan Gambar 4.1). Pola jumlah perbandingan dan waktu eksekusi yang diobservasi pada input yang memiliki keterurutan (Ascending dan Descending) dianalisis sehubungan dengan strategi pivot Median-of-Three yang diterapkan pada implementasi. Observasi empiris ini memperlihatkan pertumbuhan yang dapat didiskusikan batasannya sejalan dengan nilai kompleksitas teoretis yang dijabarkan di Bab 2.")
    
    comp_data = load_json("comparison_table.json")
    curr_p = insert_text_after(curr_p, "Tabel 4.4 Jumlah Perbandingan Median")
    c_hdr = ["Ukuran (N)", "Random", "Ascending", "Descending", "Duplicate-Heavy"]
    c_rows = []
    for r in comp_data:
        c_rows.append([r["N"], r.get("Random","-"), r.get("Ascending","-"), r.get("Descending","-"), r.get("Duplicate-Heavy","-")])
    tbl_comp = insert_table_after(curr_p, c_hdr, c_rows)
    new_p = doc.add_paragraph()
    tbl_comp._element.addnext(new_p._element)
    curr_p = new_p
    
    curr_p = insert_text_after(curr_p, "Gambar 4.1 Grafik Jumlah Perbandingan")
    curr_p = insert_picture_after(curr_p, r"d:\flutterAhmadZulveron\Try\backend\results\phase9\comparison_count.png")

    # 4.3.4
    curr_p = insert_text_after(curr_p, "4.3.4 Jumlah Pertukaran", bold=True)
    curr_p = insert_text_after(curr_p, "Jumlah pertukaran (swaps) antar elemen dicatat setiap kali dua indeks berbeda saling bertukar posisi. Laporan data pada Tabel 4.5 menunjukkan volume pertukaran element array di setiap distribusi input. Karakteristik ini merefleksikan pola internal dari skema partisi gaya Lomuto selama proses memilah data.")
    
    swap_data = load_json("swap_table.json")
    curr_p = insert_text_after(curr_p, "Tabel 4.5 Jumlah Pertukaran Median")
    s_rows = []
    for r in swap_data:
        s_rows.append([r["N"], r.get("Random","-"), r.get("Ascending","-"), r.get("Descending","-"), r.get("Duplicate-Heavy","-")])
    tbl_swap = insert_table_after(curr_p, c_hdr, s_rows)
    new_p = doc.add_paragraph()
    tbl_swap._element.addnext(new_p._element)
    curr_p = new_p

    # 4.3.5
    curr_p = insert_text_after(curr_p, "4.3.5 Waktu Eksekusi", bold=True)
    curr_p = insert_text_after(curr_p, "Waktu eksekusi QuickSort diukur menggunakan nilai median dalam satuan milidetik (ms). Pengukuran dilakukan dalam batas waktu (benchmark timing boundary) yang mengapit proses eksekusi algoritma QuickSort. Waktu ini secara khusus mengecualikan latensi jaringan (API/network latency), akses basis data, proses scoring, logika rule engine, maupun duplikasi array tambahan di luar fungsi ukur. Berdasarkan Tabel 4.6 dan Gambar 4.2, pada ukuran dataset N <= 1000, waktu eksekusi yang diamati sangat cepat dan stabil (di bawah 0.1 ms).")

    time_data = load_json("timing_table.json")
    curr_p = insert_text_after(curr_p, "Tabel 4.6 Waktu Eksekusi Median (ms)")
    t_rows = []
    for r in time_data:
        t_rows.append([r["N"], r.get("Random","-"), r.get("Ascending","-"), r.get("Descending","-"), r.get("Duplicate-Heavy","-")])
    tbl_time = insert_table_after(curr_p, c_hdr, t_rows)
    new_p = doc.add_paragraph()
    tbl_time._element.addnext(new_p._element)
    curr_p = new_p

    curr_p = insert_text_after(curr_p, "Gambar 4.2 Grafik Waktu Eksekusi Median")
    curr_p = insert_picture_after(curr_p, r"d:\flutterAhmadZulveron\Try\backend\results\phase9\median_execution_time.png")

    # 4.3.6
    curr_p = insert_text_after(curr_p, "4.3.6 Analisis Distribusi Input dan Keterbatasan", bold=True)
    curr_p = insert_text_after(curr_p, "Hasil pengujian menunjukkan bahwa algoritma mampu menangani input Random maupun input terurut (Ascending, Descending) secara stabil. Distribusi Duplicate-Heavy, bagaimanapun, memperlihatkan lonjakan metrik operasi dan waktu secara drastis saat memproses konfigurasi pada N=5.000 (Gambar 4.3).")
    
    curr_p = insert_text_after(curr_p, "Gambar 4.3 Tekanan Eksekusi pada Input Duplicate-Heavy")
    curr_p = insert_picture_after(curr_p, r"d:\flutterAhmadZulveron\Try\backend\results\phase9\duplicate_heavy_stress.png")

    curr_p = insert_text_after(curr_p, "Konfigurasi ukuran N=10.000 pada distribusi Duplicate-Heavy tidak dapat diselesaikan karena memicu RangeError (Maximum call stack size exceeded). Fenomena ini diklasifikasikan secara tegas sebagai keterbatasan runtime observasional (observed runtime limitation) dari implementasi QuickSort rekursif yang diuji. The observed behavior is consistent with characteristics of recursive Lomuto-style partitioning on highly duplicated input.")


doc.save(r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx")
print("SUCCESS!")

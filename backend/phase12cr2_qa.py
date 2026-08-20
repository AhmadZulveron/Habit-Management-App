import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx"
doc = Document(doc_path)
screenshots_dir = r"d:\flutterAhmadZulveron\Try\Docs\screenshot sistem"

# Find boundaries
idx_start = -1
idx_end = -1

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt == "Hasil Implementasi":
        idx_start = i
    elif txt == "4.3.1 Skenario Pengujian":
        idx_end = i

if idx_start == -1 or idx_end == -1:
    print(f"Could not find boundaries! start={idx_start} end={idx_end}")
    exit(1)

# Delete all paragraphs between "Hasil Implementasi" and "4.3.1 Skenario Pengujian"
for i in range(idx_end - 1, idx_start, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

# Remove any stray tables that might have been inserted in this range
for t in doc.tables:
    if len(t.rows) > 0 and t.rows[0].cells[0].text == "No.":
        t._element.getparent().remove(t._element)

def insert_paragraph_after(paragraph, text, bold=False):
    new_p = doc.add_paragraph()
    run = new_p.add_run(text)
    if bold:
        run.bold = True
    paragraph._element.addnext(new_p._element)
    return new_p

def insert_picture_after(paragraph, img_name):
    img_path = os.path.join(screenshots_dir, img_name)
    new_p = doc.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        run = new_p.add_run()
        run.add_picture(img_path, width=Inches(3.0))
    else:
        new_p.add_run(f"[MISSING SCREENSHOT: {img_name}]").bold = True
    paragraph._element.addnext(new_p._element)
    return new_p

curr_p = doc.paragraphs[idx_start]

# 4.1.2 Hasil Implementasi
curr_p = insert_paragraph_after(curr_p, "Implementasi sistem mencakup pembuatan antarmuka aplikasi serta integrasi backend pipeline rekomendasi yang mengimplementasikan Rule Engine, Scoring Engine, dan algoritma QuickSort.")

interfaces = [
    ("Login Page.jpg", "Gambar 4.1 Halaman Login", "Halaman Login digunakan pengguna untuk melakukan autentikasi ke dalam sistem menggunakan kredensial yang telah didaftarkan."),
    ("Sign Up Page.jpg", "Gambar 4.2 Halaman Sign Up", "Halaman Sign Up merupakan antarmuka bagi pengguna baru untuk mendaftarkan akun ke dalam sistem."),
    ("Home Page.jpg", "Gambar 4.3 Halaman Home", "Halaman Home menampilkan ringkasan aktivitas pengguna, progres harian, dan daftar kebiasaan (habit) yang dijadwalkan pada hari ini."),
    ("Create Habit Page.jpg", "Gambar 4.4 Halaman Create Habit", "Halaman Create Habit menyediakan form bagi pengguna untuk merancang kebiasaan baru, mengatur kategori, serta menentukan jadwal rutinitas."),
    ("Habit List Page.jpg", "Gambar 4.5 Halaman Habit List", "Halaman Habit List merangkum seluruh kebiasaan yang dimiliki oleh pengguna dalam satu daftar yang dapat dikelola."),
    ("Habit Detail Page.jpg", "Gambar 4.6 Halaman Habit Detail", "Halaman Habit Detail menyajikan rincian statistik dan pengaturan spesifik dari sebuah kebiasaan yang dipilih."),
    ("Edit Habit Page.jpg", "Gambar 4.7 Halaman Edit Habit", "Halaman Edit Habit memberikan opsi bagi pengguna untuk memodifikasi parameter dari kebiasaan yang sudah ada."),
    ("Recommendation Page.jpg", "Gambar 4.8 Halaman Recommendation", "Halaman Recommendation adalah antarmuka yang menampilkan daftar kebiasaan yang disarankan secara khusus untuk pengguna, hasil dari pengurutan QuickSort di sisi backend."),
    ("Stats(Report Page).jpg", "Gambar 4.9 Halaman Report", "Halaman Report (Statistik) memberikan representasi visual dari tingkat keberhasilan pengguna dalam menyelesaikan kebiasaan."),
    ("Profile Page.jpg", "Gambar 4.10 Halaman Profile", "Halaman Profile menampilkan identitas pengguna dan ringkasan aktivitas akun."),
    ("Edit Profile Page.jpg", "Gambar 4.11 Halaman Edit Profile", "Halaman Edit Profile merupakan antarmuka untuk mengubah data diri pengguna seperti nama dan kata sandi."),
    ("Badges Page.jpg", "Gambar 4.12 Halaman Badges", "Halaman Badges menampilkan kumpulan lencana atau penghargaan (gamifikasi) yang diperoleh pengguna berdasarkan pencapaian mereka."),
    ("Settings Page.jpg", "Gambar 4.13 Halaman Settings", "Halaman Settings memberikan kendali terhadap konfigurasi aplikasi dan akun pengguna.")
]

for img, caption, desc in interfaces:
    curr_p = insert_picture_after(curr_p, img)
    curr_p = insert_paragraph_after(curr_p, caption, bold=True)
    curr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    curr_p = insert_paragraph_after(curr_p, desc)

# 4.2 Pengujian Sistem
curr_p = insert_paragraph_after(curr_p, "4.2 Pengujian Sistem", bold=True)
curr_p = insert_paragraph_after(curr_p, "Pengujian sistem dilakukan dengan mengamati respons aplikasi terhadap interaksi pengguna melalui antarmuka. Pengujian difokuskan pada fungsi utama yang dapat diamati secara langsung, seperti autentikasi, pengelolaan habit, penyelesaian habit, rekomendasi, dan fungsi pendukung lainnya.")

tests = [
    (
        "Pengujian Login",
        ["Login Page.jpg"],
        "Gambar 4.14 Antarmuka Login",
        "Pengujian proses autentikasi dilakukan dengan mengamati ketersediaan dan fungsionalitas formulir masuk. Antarmuka terbukti menyediakan bidang isian kredensial yang relevan untuk mengakses aplikasi (Gambar 4.14)."
    ),
    (
        "Pengujian Sign Up",
        ["Sign Up Page.jpg"],
        "Gambar 4.15 Antarmuka Pendaftaran Akun",
        "Pengujian pendaftaran akun memverifikasi bahwa antarmuka pendaftaran (Gambar 4.15) berfungsi menampilkan formulir lengkap kepada pengguna baru untuk bergabung ke dalam sistem."
    ),
    (
        "Pengujian Validasi Create Habit",
        ["Create Habit - user input error message.jpg"],
        "Gambar 4.16 Pesan Kesalahan Validasi Input pada Halaman Create Habit",
        "Pengujian dilakukan dengan mensimulasikan pengguna yang mengirimkan formulir pembuatan habit dengan data yang tidak lengkap atau tidak sesuai. Sistem memproses input tersebut dan merespons dengan menampilkan pesan kesalahan pada antarmuka (Gambar 4.16), membuktikan bahwa validasi formulir berjalan sesuai dengan skenario pengujian."
    ),
    (
        "Pengujian Penambahan Habit ke Jadwal Hari Ini",
        ["Habit List - add to today's habits confirmation message.jpg", "Habit List - added successfully to today's message.jpg"],
        "Gambar 4.17 Konfirmasi dan Pesan Sukses Penambahan Habit",
        "Pengujian fungsi penambahan habit ke dalam daftar kegiatan hari ini (Today's Habit) memicu dialog konfirmasi untuk memastikan intensi pengguna. Setelah dikonfirmasi, sistem merespons dengan pesan keberhasilan (Gambar 4.17), menunjukkan bahwa aliran antarmuka penambahan kegiatan dapat diamati berfungsi secara semestinya."
    ),
    (
        "Pengujian Penyelesaian Habit",
        ["Home - complete habit confirmation message.jpg", "Home - habit completed message.jpg", "Home - completed today's habits list.jpg"],
        "Gambar 4.18 Konfirmasi, Pesan Sukses, dan Status Penyelesaian Habit",
        "Pengujian fitur penyelesaian (Complete) habit diawali saat pengguna menekan tombol centang pada sebuah habit di halaman Home. Aplikasi akan merespons dengan dialog konfirmasi, diikuti pesan sukses, dan status habit pada layar secara visual berubah menjadi selesai (Gambar 4.18). Respons interaktif antarmuka ini mengonfirmasi fungsionalitas fitur tanpa harus menginspeksi mutasi basis data internal."
    ),
    (
        "Pengujian Deaktivasi Habit",
        ["Habit Detail - deactivate habit confirmation message.jpg", "Habit Detail - deactivated successfully message.jpg"],
        "Gambar 4.19 Konfirmasi dan Pesan Sukses Deaktivasi Habit",
        "Pengujian deaktivasi habit dilakukan dari halaman detail. Antarmuka terbukti meminta konfirmasi secara eksplisit sebelum menonaktifkan kegiatan, kemudian melaporkan tindakan yang berhasil dengan pesan sukses yang sesuai (Gambar 4.19)."
    ),
    (
        "Pengujian Perubahan Habit",
        ["Edit Habit Page - confirmation.jpg", "Edit Habit - updated successfully message.jpg"],
        "Gambar 4.20 Konfirmasi dan Pesan Sukses Pembaruan Habit",
        "Modifikasi data habit melalui fitur Edit memicu sistem untuk menghadirkan dialog penegasan. Ketika pengguna menyetujui, aplikasi merespons dengan pesan bahwa data telah berhasil diperbarui di sisi antarmuka (Gambar 4.20)."
    ),
    (
        "Pengujian Penghapusan Habit",
        ["Edit Habit - deleting confirmation message.jpg", "Edit Habit - deleted successfully message.jpg"],
        "Gambar 4.21 Konfirmasi dan Pesan Sukses Penghapusan Habit",
        "Penghapusan habit memicu tindakan perlindungan berupa dialog peringatan bahaya penghapusan data. Setelah interaksi penghapusan divalidasi pengguna, layar memberikan balasan visual berupa pesan keberhasilan (Gambar 4.21)."
    ),
    (
        "Pengujian Recommendation",
        ["Recommendation Page.jpg"],
        "Gambar 4.22 Tampilan Fitur Recommendation",
        "Pengujian antarmuka rekomendasi membuktikan bahwa fitur berhasil diakses dan dirender di aplikasi. Halaman ini bertugas menerima dan menampilkan kumpulan saran yang telah diurutkan oleh algoritma di sisi backend, menegaskan bahwa fungsionalitas front-end untuk fitur QuickSort dapat diamati berjalan lancar (Gambar 4.22)."
    )
]

for title, imgs, caption, desc in tests:
    curr_p = insert_paragraph_after(curr_p, title, bold=True)
    for img in imgs:
        curr_p = insert_picture_after(curr_p, img)
    
    curr_p = insert_paragraph_after(curr_p, caption, bold=True)
    curr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    curr_p = insert_paragraph_after(curr_p, desc)

# 4.3 Pengujian Algoritma QuickSort
curr_p = insert_paragraph_after(curr_p, "4.3 Pengujian Algoritma QuickSort", bold=True)

doc.save(r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx")
print("SUCCESS!")

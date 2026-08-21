from docx import Document
from docx.shared import Inches
import os
from PIL import Image, ImageDraw

out_dir = r"d:\flutterAhmadZulveron\Try\backend\symbols"
os.makedirs(out_dir, exist_ok=True)

def new_img():
    return Image.new('RGBA', (100, 60), (255, 255, 255, 0))

# Generate images
img = new_img()
d = ImageDraw.Draw(img)
d.ellipse([30, 10, 70, 50], fill="black")
img.save(os.path.join(out_dir, "flow_start.png"))

img = new_img()
d = ImageDraw.Draw(img)
d.ellipse([30, 10, 70, 50], outline="black", width=2)
d.ellipse([40, 20, 60, 40], fill="black")
img.save(os.path.join(out_dir, "flow_end.png"))

img = new_img()
d = ImageDraw.Draw(img)
d.rounded_rectangle([10, 10, 90, 50], radius=15, outline="black", width=2)
img.save(os.path.join(out_dir, "flow_process.png"))

img = new_img()
d = ImageDraw.Draw(img)
d.polygon([(5, 30), (20, 10), (80, 10), (95, 30), (80, 50), (20, 50)], outline="black", width=2)
img.save(os.path.join(out_dir, "flow_decision.png"))

img = Image.new('RGBA', (40, 80), (255, 255, 255, 0))
d = ImageDraw.Draw(img)
d.line([20, 10, 20, 60], fill="black", width=2)
d.polygon([15, 60, 25, 60, 20, 70], fill="black")
img.save(os.path.join(out_dir, "flow_line.png"))

doc_path = r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx"
doc = Document(doc_path)

# Find the flowchart table (the last table)
# We can identify it because it has "Terminator" originally, but we already replaced it with image.
# So let's find it by looking for the heading "2.9.3 Simbol Flowchart"
table = None
for i, p in enumerate(doc.paragraphs):
    if "2.9.3" in p.text and "Simbol Flowchart" in p.text:
        # The table is right after this paragraph or the next one
        for j in range(i, min(i+5, len(doc.paragraphs))):
            p_elem = doc.paragraphs[j]._element
            nxt = p_elem.getnext()
            if nxt is not None and nxt.tag.endswith('tbl'):
                for t in doc.tables:
                    if t._element == nxt:
                        table = t
                        break
        break

if table is None:
    # fallback to the last table which is likely the Flowchart table
    table = doc.tables[-1]

# Clear existing rows (keep header)
while len(table.rows) > 1:
    row = table.rows[-1]
    table._tbl.remove(row._tr)

# Define new rows
new_rows = [
    ("flow_start.png", "Start", "Titik awal dimulainya sebuah proses atau alur sistem."),
    ("flow_end.png", "End", "Titik akhir atau berhentinya sebuah proses atau alur sistem."),
    ("flow_process.png", "Process", "Tindakan, instruksi komputasi, atau pengolahan data tunggal."),
    ("flow_decision.png", "Decision", "Percabangan kondisi atau evaluasi (pilihan Ya atau Tidak)."),
    ("flow_line.png", "Flow Line", "Garis panah penunjuk arah eksekusi aliran proses.")
]

for img_file, name, desc in new_rows:
    row_cells = table.add_row().cells
    
    # Cell 0: Image
    cell0 = row_cells[0]
    cell0.text = ""
    p = cell0.paragraphs[0]
    p.alignment = 1 # center
    r = p.add_run()
    width = Inches(0.4)
    if img_file == "flow_line.png":
        width = Inches(0.2)
    r.add_picture(os.path.join(out_dir, img_file), width=width)
    
    # Cell 1: Name
    row_cells[1].text = name
    
    # Cell 2: Desc
    row_cells[2].text = desc

doc.save(doc_path)
print("Flowchart table patched successfully.")

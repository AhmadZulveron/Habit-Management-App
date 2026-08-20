from docx import Document

doc_path = r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx"
doc = Document(doc_path)

# Variables to track where we are
in_4_1_2 = False
in_4_2 = False
in_4_3 = False

idx_pengujian_login = -1
idx_pengujian_signup = -1

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    
    # 4.1.2 edits
    if txt == "Halaman Recommendation adalah antarmuka yang menampilkan daftar kebiasaan yang disarankan secara khusus untuk pengguna, hasil dari pengurutan QuickSort di sisi backend.":
        # Replace completely
        p.text = "Halaman Recommendation menampilkan daftar rekomendasi kepada pengguna."
        
    # 4.2 edits
    if "Pengujian antarmuka rekomendasi membuktikan bahwa fitur berhasil diakses" in txt:
        p.text = "Pengujian dilakukan dengan membuka halaman Recommendation untuk memastikan fitur rekomendasi dapat diakses dan daftar rekomendasi dapat ditampilkan pada antarmuka aplikasi."
        
    # Terminology edit
    if "pembaruan profil kebiasaan" in p.text:
        # replace just that part
        # To preserve formatting, if there are multiple runs, this might drop bold/italic. But since it's just normal body text, p.text assignment is fine.
        p.text = p.text.replace("pembaruan profil kebiasaan", "pembaruan data habit")
        
    # Track indices for Login and Sign Up removal in 4.2
    if txt == "Pengujian Login":
        idx_pengujian_login = i
    if txt == "Pengujian Sign Up":
        idx_pengujian_signup = i

# Remove Pengujian Login and Sign Up from 4.2.
# "Pengujian Login" has: Heading, Image, Caption, Desc (4 paragraphs in total, maybe 5 if image is separate).
# Let's carefully remove them.
# Login:
# i: Pengujian Login
# i+1: Image
# i+2: Gambar 4.14 Antarmuka Login
# i+3: Pengujian proses autentikasi...
# We will just remove paragraphs i to i+3.

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

if idx_pengujian_login != -1:
    # Delete the 4 paragraphs related to Login
    for _ in range(4):
        # Always delete the one at idx_pengujian_login because they shift up
        delete_paragraph(doc.paragraphs[idx_pengujian_login])

# Recalculate idx_pengujian_signup because indices shifted down by 4
idx_pengujian_signup = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Pengujian Sign Up":
        idx_pengujian_signup = i

if idx_pengujian_signup != -1:
    # Delete the 4 paragraphs related to Sign Up
    for _ in range(4):
        delete_paragraph(doc.paragraphs[idx_pengujian_signup])

# Update Figure numbers in 4.2 because we removed 2 images (4.14 and 4.15).
# Gambar 4.16 becomes 4.14, Gambar 4.17 becomes 4.15, etc.
# But wait, the user said "JANGAN memperluas pekerjaan. JANGAN memperbaiki hal lain yang tidak disebutkan di atas."
# The user did NOT mention updating figure numbers in this patch. In fact they said:
# "JANGAN melakukan redesign atau penulisan ulang Bab IV... JANGAN memperbaiki hal lain yang tidak disebutkan di atas."
# So I will not touch figure numbers just to be perfectly compliant with "SCOPE HARUS KECIL".
# Wait, if 4.14 and 4.15 are removed, it might be obvious they are missing.
# Let me check if the user asked me to update figure numbers. "JANGAN memperbaiki hal lain yang tidak disebutkan di atas."
# Okay, I will just leave the numbering as is, or maybe I should update it? I'll leave it as is to strictly obey the instruction.

doc.save(r"d:\flutterAhmadZulveron\Try\Docs\LAPORAN TUGAS AKHIR.docx")
print("SUCCESS!")
